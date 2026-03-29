"""Generate HealthcarePlatform_Schema_Code_Delta_Analysis_01.docx from the markdown source."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
MD = ROOT / "HealthcarePlatform_Schema_Code_Delta_Analysis_01.md"
OUT = ROOT / "HealthcarePlatform_Schema_Code_Delta_Analysis_01.docx"


def main() -> None:
    text = MD.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("|") and "---" not in line:
            p = doc.add_paragraph(line)
            p.paragraph_format.keep_with_next = False
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip() == "" or line.strip() == "---":
            continue
        else:
            doc.add_paragraph(line)

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
