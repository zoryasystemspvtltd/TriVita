"""Generate HealthcarePlatform_Technical_Architecture_and_Documentation.docx from markdown."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
MD = ROOT / "HealthcarePlatform_Technical_Architecture_and_Documentation.md"
OUT = ROOT / "HealthcarePlatform_Technical_Architecture_and_Documentation.docx"


def main() -> None:
    text = MD.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_mermaid = False
    mermaid_buf: list[str] = []

    for line in text.splitlines():
        if line.strip() == "```mermaid":
            in_mermaid = True
            mermaid_buf = []
            continue
        if in_mermaid:
            if line.strip() == "```":
                doc.add_paragraph("Mermaid diagram (render in Markdown viewer or paste into mermaid.live):", style="Intense Quote")
                p = doc.add_paragraph("\n".join(mermaid_buf))
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                in_mermaid = False
                mermaid_buf = []
            else:
                mermaid_buf.append(line)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("|") and "---" not in line:
            doc.add_paragraph(line)
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip() == "" or line.strip() == "---":
            continue
        elif line.strip().startswith("```"):
            continue
        else:
            doc.add_paragraph(line)

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
